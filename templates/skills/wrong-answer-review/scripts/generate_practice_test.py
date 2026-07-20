#!/usr/bin/env python3
"""
错题练习卷 Word 文档生成脚本

从 JSON 数据文件生成格式化的 Word 练习卷，包含：
- 薄弱环节分析
- 原试卷错题
- 变式强化练习
- 参考答案与解析

用法：
    python3 generate_practice_test.py --data data.json --output output.docx
"""

import json
import argparse
import sys
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("错误：需要安装 python-docx 库")
    print("请执行：pip install python-docx")
    sys.exit(1)


# ==================== 样式配置 ====================

# 颜色定义
COLOR_PRIMARY = RGBColor(0x1A, 0x56, 0xDB)      # 主色调 - 深蓝
COLOR_SECONDARY = RGBColor(0x2E, 0x7D, 0x32)     # 辅助色 - 深绿
COLOR_ACCENT = RGBColor(0xE6, 0x51, 0x00)         # 强调色 - 橙红
COLOR_TEXT = RGBColor(0x33, 0x33, 0x33)            # 正文色 - 深灰
COLOR_LIGHT_TEXT = RGBColor(0x66, 0x66, 0x66)      # 浅文字
COLOR_BG_LIGHT = RGBColor(0xF5, 0xF5, 0xF5)       # 浅背景
COLOR_BG_HEADER = RGBColor(0xE3, 0xF2, 0xFD)      # 表头背景 - 浅蓝
COLOR_BG_WRONG = RGBColor(0xFF, 0xEB, 0xEE)        # 错题背景 - 浅红
COLOR_BORDER = RGBColor(0xBD, 0xBD, 0xBD)          # 边框色


def set_cell_shading(cell, color_hex):
    """设置单元格背景色"""
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_paragraph_spacing(paragraph, before=0, after=0):
    """设置段落间距"""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)


def create_title_page(doc, student_info):
    """创建标题页"""
    # 添加空行
    for _ in range(6):
        doc.add_paragraph()

    # 主标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("错题练习卷")
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = COLOR_PRIMARY
    add_paragraph_spacing(title, before=0, after=12)

    # 副标题 - 年级学科
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    grade = student_info.get("grade", "")
    subject = student_info.get("subject", "")
    textbook = student_info.get("textbook", "")
    grade_info = f"{grade} · {subject}"
    if textbook:
        grade_info += f" · {textbook}"
    run = subtitle.add_run(grade_info)
    run.font.size = Pt(18)
    run.font.color.rgb = COLOR_LIGHT_TEXT
    add_paragraph_spacing(subtitle, before=0, after=24)

    # 日期
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    today = datetime.now().strftime("%Y年%m月%d日")
    run = date_para.add_run(f"生成日期：{today}")
    run.font.size = Pt(12)
    run.font.color.rgb = COLOR_LIGHT_TEXT

    # 分页
    doc.add_page_break()


def create_weakness_analysis(doc, analysis):
    """创建薄弱环节分析部分"""
    # 章节标题
    heading = doc.add_heading("薄弱环节分析", level=1)
    for run in heading.runs:
        run.font.color.rgb = COLOR_PRIMARY

    # 分析摘要
    if analysis.get("summary"):
        summary_para = doc.add_paragraph()
        run = summary_para.add_run(analysis["summary"])
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_TEXT
        add_paragraph_spacing(summary_para, before=6, after=12)

    # 题型统计表
    if analysis.get("by_question_type"):
        type_heading = doc.add_heading("各题型错误统计", level=2)
        for run in type_heading.runs:
            run.font.color.rgb = COLOR_SECONDARY

        # 创建表格
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 设置列宽
        widths = [Cm(4), Cm(3), Cm(3), Cm(6)]
        for i, width in enumerate(widths):
            table.columns[i].width = width

        # 表头
        headers = ["题型", "错题数", "占比", "涉及知识点"]
        for i, header_text in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(header_text)
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = COLOR_TEXT
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_shading(cell, "E3F2FD")

        # 数据行
        for item in analysis["by_question_type"]:
            row = table.add_row()
            
            # 题型
            cell0 = row.cells[0]
            cell0.text = ""
            p = cell0.paragraphs[0]
            run = p.add_run(item.get("type", ""))
            run.font.size = Pt(10)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 错题数
            cell1 = row.cells[1]
            cell1.text = ""
            p = cell1.paragraphs[0]
            run = p.add_run(str(item.get("count", 0)))
            run.font.size = Pt(10)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 占比
            cell2 = row.cells[2]
            cell2.text = ""
            p = cell2.paragraphs[0]
            run = p.add_run(item.get("percentage", ""))
            run.font.size = Pt(10)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 知识点
            cell3 = row.cells[3]
            cell3.text = ""
            p = cell3.paragraphs[0]
            kp_text = "、".join(item.get("knowledge_points", []))
            run = p.add_run(kp_text)
            run.font.size = Pt(10)

        doc.add_paragraph()  # 空行

    # 薄弱知识点列表
    if analysis.get("knowledge_points"):
        kp_heading = doc.add_heading("薄弱知识点", level=2)
        for run in kp_heading.runs:
            run.font.color.rgb = COLOR_SECONDARY

        kp_para = doc.add_paragraph()
        kp_text = "  ·  ".join(analysis["knowledge_points"])
        run = kp_para.add_run(kp_text)
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_TEXT
        add_paragraph_spacing(kp_para, before=6, after=12)

    # 复习建议
    if analysis.get("suggestions"):
        sug_heading = doc.add_heading("复习建议", level=2)
        for run in sug_heading.runs:
            run.font.color.rgb = COLOR_SECONDARY

        for i, suggestion in enumerate(analysis["suggestions"], 1):
            sug_para = doc.add_paragraph()
            run = sug_para.add_run(f"{i}. {suggestion}")
            run.font.size = Pt(11)
            run.font.color.rgb = COLOR_TEXT
            add_paragraph_spacing(sug_para, before=3, after=3)

    doc.add_page_break()


def create_wrong_questions_section(doc, wrong_questions):
    """创建原试卷错题部分"""
    heading = doc.add_heading("第一部分：原试卷错题", level=1)
    for run in heading.runs:
        run.font.color.rgb = COLOR_ACCENT

    if not wrong_questions:
        empty_para = doc.add_paragraph()
        run = empty_para.add_run("未识别到错题。")
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_LIGHT_TEXT
        return

    # 按题型分组
    grouped = {}
    for q in wrong_questions:
        q_type = q.get("question_type", "其他")
        if q_type not in grouped:
            grouped[q_type] = []
        grouped[q_type].append(q)

    for q_type, questions in grouped.items():
        # 题型小标题
        type_heading = doc.add_heading(q_type, level=2)
        for run in type_heading.runs:
            run.font.color.rgb = COLOR_SECONDARY

        for q in questions:
            # 题目容器 - 使用浅红色背景
            q_para = doc.add_paragraph()
            q_para.paragraph_format.left_indent = Cm(0.5)
            
            # 来源标记
            source = q.get("source_image", "")
            if source:
                source_run = q_para.add_run(f"[{source}] ")
                source_run.font.size = Pt(9)
                source_run.font.color.rgb = COLOR_LIGHT_TEXT
                source_run.font.italic = True

            # 题号和内容
            q_num = q.get("question_number", "")
            q_content = q.get("content", "")
            run = q_para.add_run(f"第 {q_num} 题：{q_content}")
            run.font.size = Pt(11)
            run.font.color.rgb = COLOR_TEXT
            add_paragraph_spacing(q_para, before=6, after=3)

            # 选项
            options = q.get("options", [])
            if options:
                for opt in options:
                    opt_para = doc.add_paragraph()
                    opt_para.paragraph_format.left_indent = Cm(1.0)
                    run = opt_para.add_run(opt)
                    run.font.size = Pt(11)
                    run.font.color.rgb = COLOR_TEXT
                    add_paragraph_spacing(opt_para, before=1, after=1)

            # 学生答案和正确答案
            answer_para = doc.add_paragraph()
            answer_para.paragraph_format.left_indent = Cm(0.5)
            
            student_ans = q.get("student_answer", "")
            correct_ans = q.get("correct_answer", "")
            
            if student_ans:
                run = answer_para.add_run("学生答案：")
                run.font.size = Pt(10)
                run.font.bold = True
                run.font.color.rgb = COLOR_ACCENT
                run = answer_para.add_run(student_ans)
                run.font.size = Pt(10)
                run.font.color.rgb = COLOR_ACCENT
            
            if correct_ans:
                if student_ans:
                    run = answer_para.add_run("    ")
                run = answer_para.add_run("正确答案：")
                run.font.size = Pt(10)
                run.font.bold = True
                run.font.color.rgb = COLOR_SECONDARY
                run = answer_para.add_run(correct_ans)
                run.font.size = Pt(10)
                run.font.color.rgb = COLOR_SECONDARY

            add_paragraph_spacing(answer_para, before=3, after=9)

    doc.add_page_break()


def create_variant_questions_section(doc, variant_questions):
    """创建变式强化练习部分"""
    heading = doc.add_heading("第二部分：变式强化练习", level=1)
    for run in heading.runs:
        run.font.color.rgb = COLOR_PRIMARY

    if not variant_questions:
        empty_para = doc.add_paragraph()
        run = empty_para.add_run("未生成变式练习题。")
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_LIGHT_TEXT
        return

    # 按关联错题分组
    grouped = {}
    for vq in variant_questions:
        related = vq.get("related_wrong_question", "unknown")
        if related not in grouped:
            grouped[related] = []
        grouped[related].append(vq)

    q_index = 1
    for related_num, questions in grouped.items():
        # 关联标题
        ref_heading = doc.add_heading(f"针对第 {related_num} 题的变式练习", level=2)
        for run in ref_heading.runs:
            run.font.color.rgb = COLOR_SECONDARY

        for vq in questions:
            q_para = doc.add_paragraph()
            q_para.paragraph_format.left_indent = Cm(0.5)
            
            run = q_para.add_run(f"{q_index}. ")
            run.font.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = COLOR_TEXT
            
            run = q_para.add_run(vq.get("content", ""))
            run.font.size = Pt(11)
            run.font.color.rgb = COLOR_TEXT
            add_paragraph_spacing(q_para, before=6, after=3)

            # 选项
            options = vq.get("options", [])
            if options:
                for opt in options:
                    opt_para = doc.add_paragraph()
                    opt_para.paragraph_format.left_indent = Cm(1.0)
                    run = opt_para.add_run(opt)
                    run.font.size = Pt(11)
                    run.font.color.rgb = COLOR_TEXT
                    add_paragraph_spacing(opt_para, before=1, after=1)

            # 知识点标签
            kp = vq.get("knowledge_point", "")
            if kp:
                kp_para = doc.add_paragraph()
                kp_para.paragraph_format.left_indent = Cm(0.5)
                run = kp_para.add_run(f"知识点：{kp}")
                run.font.size = Pt(9)
                run.font.italic = True
                run.font.color.rgb = COLOR_LIGHT_TEXT
                add_paragraph_spacing(kp_para, before=1, after=9)

            q_index += 1

    doc.add_page_break()


def create_answer_key_section(doc, variant_questions):
    """创建参考答案与解析部分"""
    heading = doc.add_heading("参考答案与解析", level=1)
    for run in heading.runs:
        run.font.color.rgb = COLOR_PRIMARY

    if not variant_questions:
        empty_para = doc.add_paragraph()
        run = empty_para.add_run("无变式练习题答案。")
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_LIGHT_TEXT
        return

    for i, vq in enumerate(variant_questions, 1):
        # 题号
        ans_para = doc.add_paragraph()
        ans_para.paragraph_format.left_indent = Cm(0.5)
        
        run = ans_para.add_run(f"第 {i} 题")
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_TEXT
        add_paragraph_spacing(ans_para, before=9, after=3)

        # 答案
        answer_para = doc.add_paragraph()
        answer_para.paragraph_format.left_indent = Cm(1.0)
        run = answer_para.add_run("答案：")
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_SECONDARY
        run = answer_para.add_run(vq.get("answer", ""))
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_TEXT
        add_paragraph_spacing(answer_para, before=3, after=3)

        # 解析
        explanation = vq.get("explanation", "")
        if explanation:
            exp_para = doc.add_paragraph()
            exp_para.paragraph_format.left_indent = Cm(1.0)
            run = exp_para.add_run("解析：")
            run.font.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = COLOR_SECONDARY
            run = exp_para.add_run(explanation)
            run.font.size = Pt(11)
            run.font.color.rgb = COLOR_TEXT
            add_paragraph_spacing(exp_para, before=3, after=3)


def create_practice_test(data, output_path):
    """生成错题练习卷 Word 文档"""
    doc = Document()

    # 设置默认字体
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Microsoft YaHei"
    font.size = Pt(11)
    font.color.rgb = COLOR_TEXT

    # 设置页边距
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # 1. 标题页
    student_info = data.get("student_info", {})
    create_title_page(doc, student_info)

    # 2. 薄弱环节分析
    weakness_analysis = data.get("weakness_analysis", {})
    create_weakness_analysis(doc, weakness_analysis)

    # 3. 原试卷错题
    wrong_questions = data.get("wrong_questions", [])
    create_wrong_questions_section(doc, wrong_questions)

    # 4. 变式强化练习
    variant_questions = data.get("variant_questions", [])
    create_variant_questions_section(doc, variant_questions)

    # 5. 参考答案与解析
    create_answer_key_section(doc, variant_questions)

    # 保存文档
    doc.save(output_path)
    print(f"✅ 练习卷已生成: {output_path}")

    # 输出统计信息
    print(f"   年级学科: {student_info.get('grade', '未知')} · {student_info.get('subject', '未知')}")
    print(f"   错题数量: {len(wrong_questions)} 道")
    print(f"   变式题目: {len(variant_questions)} 道")
    
    # 统计题型分布
    type_counts = {}
    for q in wrong_questions:
        q_type = q.get("question_type", "其他")
        type_counts[q_type] = type_counts.get(q_type, 0) + 1
    if type_counts:
        type_summary = "、".join([f"{k} {v}道" for k, v in type_counts.items()])
        print(f"   题型分布: {type_summary}")


def main():
    parser = argparse.ArgumentParser(
        description="生成错题练习卷 Word 文档",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例：
  python3 generate_practice_test.py --data data.json --output 练习卷.docx

数据文件格式请参考 SKILL.md 中的 JSON 示例。
        """
    )
    parser.add_argument(
        "--data",
        required=True,
        help="JSON 数据文件路径"
    )
    parser.add_argument(
        "--output",
        required=True,
        help="输出 Word 文档路径"
    )

    args = parser.parse_args()

    # 读取数据
    try:
        with open(args.data, "r", encoding="utf-8") as f:
            data = json.load(f)
    except FileNotFoundError:
        print(f"错误：找不到数据文件 {args.data}")
        sys.exit(1)
    except json.JSONDecodeError as e:
        print(f"错误：JSON 格式不正确 - {e}")
        sys.exit(1)

    # 生成文档
    create_practice_test(data, args.output)


if __name__ == "__main__":
    main()
