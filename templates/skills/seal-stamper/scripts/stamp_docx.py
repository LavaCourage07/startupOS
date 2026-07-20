#!/usr/bin/env python3
"""
在 Word 文档中插入浮动印章图片
通过直接操作 OOXML XML 实现 <wp:anchor> 浮动图片
"""

import argparse
import sys
from pathlib import Path
from typing import List, Optional, Tuple

from docx import Document
from docx.oxml.ns import qn
from lxml import etree


# ── 常量 ──────────────────────────────────────────────
MM_TO_EMU = 36000  # 1mm = 36000 EMU

# 印章尺寸默认值（圆形公章直径 42mm）
DEFAULT_SEAL_SIZE_MM = 42

# 盖章位置关键词
SEAL_KEYWORDS = [
    "（盖章）", "(盖章)", "盖章处", "公司盖章",
    "甲方盖章", "乙方盖章", "丙方盖章",
    "盖章", "（章）", "(章)",
]

# XML 命名空间
WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
PIC_NS = "http://schemas.openxmlformats.org/drawingml/2006/picture"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

# lxml XPath 用的命名空间映射
_NS_MAP = {"a": A_NS, "pic": PIC_NS, "wp": WP_NS, "r": R_NS}


def mm_to_emu(mm: float) -> int:
    """毫米转 EMU"""
    return int(mm * MM_TO_EMU)


def find_seal_positions(doc: Document) -> List[Tuple[int, str]]:
    """
    扫描文档，找到包含盖章关键词的段落
    返回: [(段落索引, 匹配的关键词), ...]
    """
    positions = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        for keyword in SEAL_KEYWORDS:
            if keyword in text:
                positions.append((i, keyword))
                break
    return positions


def _build_anchor_xml(
    r_embed: str,
    pos_x_emu: int,
    pos_y_emu: int,
    width_emu: int,
    height_emu: int,
    behind_document: bool = False,
    doc_pr_id: int = 1,
) -> bytes:
    """构建 wp:anchor 的 XML 字节串"""
    behind = "1" if behind_document else "0"
    anchor_xml = (
        f'<wp:anchor xmlns:wp="{WP_NS}" '
        f'xmlns:a="{A_NS}" '
        f'xmlns:pic="{PIC_NS}" '
        f'xmlns:r="{R_NS}" '
        f'distT="0" distB="0" distL="0" distR="0" '
        f'simplePos="0" relativeHeight="25" '
        f'behindDoc="{behind}" locked="0" '
        f'layoutInCell="1" allowOverlap="1">'
        f'<wp:simplePos x="0" y="0"/>'
        f'<wp:positionH relativeFrom="page">'
        f'<wp:posOffset>{pos_x_emu}</wp:posOffset>'
        f'</wp:positionH>'
        f'<wp:positionV relativeFrom="paragraph">'
        f'<wp:posOffset>{pos_y_emu}</wp:posOffset>'
        f'</wp:positionV>'
        f'<wp:extent cx="{width_emu}" cy="{height_emu}"/>'
        f'<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        f'<wp:wrapNone/>'
        f'<wp:docPr id="{doc_pr_id}" name="seal_{doc_pr_id}"/>'
        f'<wp:cNvGraphicFramePr>'
        f'<a:graphicFrameLocks noChangeAspect="1"/>'
        f'</wp:cNvGraphicFramePr>'
        f'<a:graphic>'
        f'<a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        f'<pic:pic>'
        f'<pic:nvPicPr>'
        f'<pic:cNvPr id="0" name="seal.png"/>'
        f'<pic:cNvPicPr/>'
        f'</pic:nvPicPr>'
        f'<pic:blipFill>'
        f'<a:blip r:embed="{r_embed}"/>'
        f'<a:stretch><a:fillRect/></a:stretch>'
        f'</pic:blipFill>'
        f'<pic:spPr>'
        f'<a:xfrm><a:off x="0" y="0"/>'
        f'<a:ext cx="{width_emu}" cy="{height_emu}"/>'
        f'</a:xfrm>'
        f'<a:prstGeom prst="rect"><a:avLst/></a:prstGeom>'
        f'</pic:spPr>'
        f'</pic:pic>'
        f'</a:graphicData>'
        f'</a:graphic>'
        f'</wp:anchor>'
    )
    return anchor_xml.encode("utf-8")


def add_anchor_image(
    doc: Document,
    paragraph,
    image_path: str,
    pos_x_emu: int,
    pos_y_emu: int,
    width_emu: int,
    height_emu: int,
    behind_document: bool = False,
    doc_pr_id: int = 1,
) -> None:
    """
    在指定段落插入浮动（anchor）图片
    behind_document=False → 浮于文字上方
    behind_document=True  → 衬于文字下方

    实现思路：
    1. 先用 python-docx 的 add_picture 临时插入内联图片（自动注册图片资源）
    2. 从内联图片的 XML 中提取 r:embed 关系 ID
    3. 删除临时内联图片
    4. 用该关系 ID 构建 anchor XML 并插入
    """
    from docx.shared import Inches

    # Step 1: 临时插入内联图片以注册图片资源
    temp_run = paragraph.add_run()
    temp_inline_shape = temp_run.add_picture(image_path, width=Inches(1))

    # Step 2: 从 _inline 元素提取 blip 的 r:embed
    # 使用完整命名空间 URI 进行 XPath 查询（python-docx 的 qn 不支持多级路径）
    inline_elem = temp_inline_shape._inline
    blip = inline_elem.find(f'.//{{{A_NS}}}blip')
    if blip is None:
        raise RuntimeError("无法从内联图片中提取 blip 关系 ID")
    r_embed = blip.get(f'{{{R_NS}}}embed')

    # Step 3: 移除临时插入的 run
    paragraph._element.remove(temp_run._element)

    # Step 4: 构建 anchor 元素
    anchor_bytes = _build_anchor_xml(
        r_embed, pos_x_emu, pos_y_emu,
        width_emu, height_emu,
        behind_document, doc_pr_id,
    )
    anchor_elem = etree.fromstring(anchor_bytes)

    # Step 5: 将 anchor 包装到 w:drawing > w:r 中并追加到段落
    new_run = etree.SubElement(paragraph._element, qn("w:r"))
    new_drawing = etree.SubElement(new_run, qn("w:drawing"))
    new_drawing.append(anchor_elem)


def stamp_document(
    docx_path: str,
    seal_image_path: str,
    output_path: str,
    seal_size_mm: float = DEFAULT_SEAL_SIZE_MM,
    pos_x_mm: Optional[float] = None,
    pos_y_mm: Optional[float] = None,
    paragraph_index: Optional[int] = None,
    behind_document: bool = False,
    auto_detect: bool = True,
) -> str:
    """
    主函数：在 Word 文档中插入浮动印章

    参数:
        docx_path: Word 文档路径
        seal_image_path: 去背景后的印章图片路径
        output_path: 输出文件路径
        seal_size_mm: 印章直径（毫米）
        pos_x_mm: 水平位置（毫米），相对于页面左边距
        pos_y_mm: 垂直位置（毫米），相对于段落顶部
        paragraph_index: 目标段落索引
        behind_document: True=衬于文字下方, False=浮于文字上方
        auto_detect: 是否自动检测盖章位置

    返回:
        输出文件路径
    """
    doc = Document(docx_path)

    width_emu = mm_to_emu(seal_size_mm)
    height_emu = mm_to_emu(seal_size_mm)
    section = doc.sections[0]

    doc_pr_id = 1  # 每个图片需要唯一的 id

    if auto_detect and paragraph_index is None:
        positions = find_seal_positions(doc)
        if not positions:
            print("未找到自动盖章位置，请在文档中使用'（盖章）'等标记，或手动指定段落索引")
            sys.exit(1)

        for idx, keyword in positions:
            paragraph = doc.paragraphs[idx]
            print(f"在段落 {idx} 处盖章（匹配关键词: {keyword}）")

            # 印章放在右侧，距离右边距约 20mm
            page_width_emu = section.page_width
            right_margin_emu = section.right_margin
            pos_x_emu = page_width_emu - right_margin_emu - mm_to_emu(seal_size_mm) - mm_to_emu(20)
            pos_y_emu = mm_to_emu(5)

            add_anchor_image(
                doc, paragraph, seal_image_path,
                pos_x_emu, pos_y_emu,
                width_emu, height_emu,
                behind_document, doc_pr_id,
            )
            doc_pr_id += 1
    else:
        if paragraph_index is None:
            paragraph_index = len(doc.paragraphs) - 1
            print(f"未指定段落，默认在最后一个段落（索引 {paragraph_index}）盖章")

        paragraph = doc.paragraphs[paragraph_index]

        if pos_x_mm is not None and pos_y_mm is not None:
            pos_x_emu = mm_to_emu(pos_x_mm)
            pos_y_emu = mm_to_emu(pos_y_mm)
        else:
            page_width_emu = section.page_width
            right_margin_emu = section.right_margin
            pos_x_emu = page_width_emu - right_margin_emu - mm_to_emu(seal_size_mm) - mm_to_emu(20)
            pos_y_emu = mm_to_emu(5)

        add_anchor_image(
            doc, paragraph, seal_image_path,
            pos_x_emu, pos_y_emu,
            width_emu, height_emu,
            behind_document, doc_pr_id,
        )

    doc.save(output_path)
    print(f"盖章完成: {output_path}")
    return output_path


def main():
    parser = argparse.ArgumentParser(description="Word 文档盖章")
    parser.add_argument("docx", help="Word 文档路径")
    parser.add_argument("seal", help="去背景后的印章图片路径")
    parser.add_argument("output", help="输出文件路径")
    parser.add_argument("--size", type=float, default=DEFAULT_SEAL_SIZE_MM,
                        help=f"印章直径（毫米），默认 {DEFAULT_SEAL_SIZE_MM}")
    parser.add_argument("--pos-x", type=float, default=None,
                        help="水平位置（毫米，相对于页面左边距）")
    parser.add_argument("--pos-y", type=float, default=None,
                        help="垂直位置（毫米，相对于段落顶部）")
    parser.add_argument("--paragraph", type=int, default=None,
                        help="目标段落索引")
    parser.add_argument("--behind", action="store_true",
                        help="衬于文字下方（默认浮于文字上方）")
    parser.add_argument("--no-auto-detect", action="store_true",
                        help="禁用自动检测盖章位置")
    args = parser.parse_args()

    if not Path(args.docx).exists():
        print(f"错误: Word 文档不存在: {args.docx}")
        sys.exit(1)
    if not Path(args.seal).exists():
        print(f"错误: 印章图片不存在: {args.seal}")
        sys.exit(1)

    stamp_document(
        args.docx, args.seal, args.output,
        seal_size_mm=args.size,
        pos_x_mm=args.pos_x,
        pos_y_mm=args.pos_y,
        paragraph_index=args.paragraph,
        behind_document=args.behind,
        auto_detect=not args.no_auto_detect,
    )


if __name__ == "__main__":
    main()
